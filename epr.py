import streamlit as st
import pandas as pd
import io
import os
from docx import Document
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Configuração da Página
st.set_page_config(page_title="Simulador de ERP & Gabarito Operacional", layout="wide", page_icon="⚙️")

# --- BANCO DE DADOS DIDÁTICO DE ERPS ---
MAPEAMENTO_ERP = {
    "SAP S/4HANA": {
        "AV-1001": {"transacao": "ME21N / MIGO", "modulo": "MM (Materiais)", "acao": "Entrada de Nota e Pedido de Compra de Chapa de Aço."},
        "AV-1002": {"transacao": "CR02 / PP01", "modulo": "PP (Produção)", "acao": "Ajuste de capacidade do centro de trabalho de Solda."},
        "AV-1003": {"transacao": "VT01N / LE-TRA", "modulo": "LES (Logística)", "acao": "Registro de avaria de frota e reatribuição de transporte."},
        "AV-1004": {"transacao": "ME28 / ME22N", "modulo": "MM (Aprovações)", "acao": "Liberação de estratégia de bloqueio do pedido ME21N."}
    },
    "TOTVS Protheus": {
        "AV-1001": {"transacao": "MATA103 / MATA120", "modulo": "Compras / Estoque", "acao": "Inclusão de Solicitação de Compra para matéria-prima."},
        "AV-1002": {"transacao": "SH6 / MATA630", "modulo": "PCP", "acao": "Reagendamento de Carga de Máquina na linha de solda."},
        "AV-1003": {"transacao": "TMSA010", "modulo": "TMS (Transporte)", "acao": "Manutenção de viagem e substituição do veículo da rota."},
        "AV-1004": {"transacao": "MATA097", "modulo": "Aprovações", "acao": "Aprovação do documento de compra bloqueado na alçada."}
    },
    "Vertis ERP": {
        "AV-1001": {"transacao": "SUP-0010", "modulo": "Suprimentos", "acao": "Emissão emergencial de ordem de fornecimento."},
        "AV-1002": {"transacao": "IND-0045", "modulo": "Chão de Fábrica", "acao": "Redistribuição de lote de produção sob gargalo."},
        "AV-1003": {"transacao": "LOG-0102", "modulo": "Expedição", "acao": "Abertura de O.S. para frota e transbordo de carga."},
        "AV-1004": {"transacao": "ADM-0080", "modulo": "Financeiro/Fiscal", "acao": "Desbloqueio de pendência documental para liberação."}
    },
    "Oracle EBS": {
        "AV-1001": {"transacao": "PO_POXPOEPO", "modulo": "Purchasing", "acao": "Criar requisição de compra de material faltante."},
        "AV-1002": {"transacao": "WIP_WIPDISPO", "modulo": "Work in Process", "acao": "Reequilíbrio de fluxo de trabalho de produção."},
        "AV-1003": {"transacao": "WSH_WSHTRLVE", "modulo": "Shipping Execution", "acao": "Redirecionamento de entrega e parada logística."},
        "AV-1004": {"transacao": "PO_POXAPPRO", "modulo": "Approvals", "acao": "Aprovação hierárquica de pedido de compra."}
    }
}

# --- FUNÇÕES DE EXPORTAÇÃO ---
def gerar_word(df, erp_selecionado):
    doc = Document()
    doc.add_heading(f'📊 Relatório de Diagnóstico & Gabarito - {erp_selecionado}', level=0)
    
    doc.add_heading('1. O que é este documento?', level=2)
    doc.add_paragraph("Este é um gabarito explicativo e simulação de organização das paradas industriais dentro do ERP selecionado.")
    
    doc.add_heading('2. Organização e Transações no ERP', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ['Código', 'Centro Trabalho', 'Transação / Tela', 'Módulo ERP', 'Ação Sugerida']
    for i, title in enumerate(headers):
        table.rows[0].cells[i].text = title
        
    for _, row in df.iterrows():
        cod = str(row.get('Codigo_Aviso', ''))
        info_erp = MAPEAMENTO_ERP[erp_selecionado].get(cod, {"transacao": "N/A", "modulo": "N/A", "acao": "Ação genérica"})
        
        row_cells = table.add_row().cells
        row_cells[0].text = cod
        row_cells[1].text = str(row.get('Centro_Trabalho', ''))
        row_cells[2].text = info_erp['transacao']
        row_cells[3].text = info_erp['modulo']
        row_cells[4].text = info_erp['acao']

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_excel(df, erp_selecionado):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Dados Originais', index=False)
        
        # Mapeando ERP no Excel
        dados_erp = []
        for _, row in df.iterrows():
            cod = str(row.get('Codigo_Aviso', ''))
            info = MAPEAMENTO_ERP[erp_selecionado].get(cod, {"transacao": "N/A", "modulo": "N/A", "acao": "N/A"})
            dados_erp.append({
                "Código": cod,
                "Centro": row.get('Centro_Trabalho', ''),
                "Módulo ERP": info['modulo'],
                "Transação/Tela": info['transacao'],
                "Procedimento": info['acao']
            })
        pd.DataFrame(dados_erp).to_excel(writer, sheet_name=f'Gabarito_{erp_selecionado}', index=False)
    buffer.seek(0)
    return buffer

# --- INTERFACE PRINCIPAL ---
st.title("🤖 Simulador de ERP & Gabarito Operacional")
st.markdown("---")

# Barra Lateral - Configurações e Botões de Troca
st.sidebar.header("🕹️ Painel de Controle")
erp_escolhido = st.sidebar.selectbox("Selecione o ERP de Origem/Destino:", list(MAPEAMENTO_ERP.keys()))
nivel_usuario = st.sidebar.radio("Nível de Aprendizado / Perfil:", ["Júnior (Passo a Passo)", "Pleno (Análise Tática)", "Sênior (Gestão Estratégica)"])

# Upload do Arquivo
st.subheader("📂 1. Caixinha de Upload do Relatório (CSV)")
uploaded_file = st.file_uploader("Arraste seu relatório CSV aqui para o robô analisar:", type=["csv"])

# Dados padrão (Caso o usuário não suba arquivo imediato)
dados_padrao = """Codigo_Aviso;Centro_Trabalho;Descricao_Gargalo;Status_Linha;Dias_Atraso;Custo_Parada_R$
AV-1001;Prensa-01;Falta de Insumo (Chapa de Aço);PARADA;4.5;18500.00
AV-1002;Solda-02;Sobrecarga de Ordens (Gargalo de Capacidade);LENTA;2.0;6200.00
AV-1003;Rota-Norte;Atraso na Entrega da Carga / Frota Quebrada;EM TRÂNSITO;3.0;9400.00
AV-1004;Montagem-01;Aguardando Liberação de Pedido ME21N;PARADA;5.0;22100.00"""

if uploaded_file is not None:
    try:
        df = pd.read_csv(uploaded_file, sep=None, engine='python')
    except Exception:
        df = pd.read_csv(io.StringIO(dados_padrao), sep=';')
else:
    st.info("💡 Exibindo dados de exemplo. Faça upload do seu CSV para substituir.")
    df = pd.read_csv(io.StringIO(dados_padrao), sep=';')

st.markdown("---")

# Visualização de Como Faria no ERP Escolhido
st.subheader(f"🗺️ 2. Gabarito Operacional: Como organizaria no **{erp_escolhido}**")

col_dados, col_gabarito = st.columns([1, 1.2])

with col_dados:
    st.markdown("**📋 Dados Brutos Lidos:**")
    st.dataframe(df, use_container_width=True)

with col_gabarito:
    st.markdown(f"**⚙️ Mapeamento de Transações no {erp_escolhido}:**")
    
    linhas_gabarito = []
    for _, row in df.iterrows():
        cod = str(row.get('Codigo_Aviso', ''))
        info = MAPEAMENTO_ERP[erp_escolhido].get(cod, {
            "transacao": "TRANS_GENERICA", 
            "modulo": "Geral", 
            "acao": "Cadastrar e tratar chamado no sistema."
        })
        linhas_gabarito.append({
            "Código": cod,
            "Módulo": info["modulo"],
            "Transação / Tela": info["transacao"],
            "Instrução": info["acao"]
        })
    df_gabarito = pd.DataFrame(linhas_gabarito)
    st.dataframe(df_gabarito, use_container_width=True)

st.markdown("---")

# --- MANUAL DIDÁTICO ADAPTATIVO POR NÍVEL ---
st.subheader(f"📘 3. Manual Didático de Execução - Perfil: {nivel_usuario}")

if "Júnior" in nivel_usuario:
    st.success("👶 **Instruções para Analista Júnior (Tudo Mastigado e Explicativo):**")
    st.markdown("""
    1. **Acesse o ERP selecionado** na sua máquina com seu usuário e senha.
    2. Digite o código da transação/tela exatamente como mostrado na tabela do gabarito acima.
    3. Para **AV-1004 (Montagem-01)**: É a prioridade maxima! Abra a tela de liberação e solicite a aprovação do seu gestor.
    4. Para **AV-1001 (Prensa-01)**: Vá no módulo de estoque e verifique se há saldo de chapa de aço em outro almoxarifado.
    5. **Dica Simples:** Não altere dados sem salvar o código da ordem de serviço.
    """)

elif "Pleno" in nivel_usuario:
    st.warning("🧑‍💻 **Instruções para Analista Pleno (Visão Tática e Processos):**")
    st.markdown("""
    1. **Análise de Causa Raiz:** Note que 72% dos custos de parada são administrativos e de suprimentos (AV-1001 e AV-1004).
    2. **Balanceamento de Capacidade:** Na linha de Solda-02, re-equilibre o plano master de produção (MPS) no ERP para eliminar a lentidão.
    3. **Logística:** Acione a manutenção de frota e verifique a apólice de seguro para transbordo da Rota-Norte.
    """)

else:
    st.error("👨‍💼 **Instruções para Analista Sênior / Gerência (Visão Estratégica):**")
    st.markdown("""
    1. **Impacto no EBTIDA:** O prejuízo acumulado de **R$ 56.200,00** exige revisão imediata nos SLAs de compras e fornecedores.
    2. **Governança de TI/ERP:** Automatizar o workflow de aprovação da transação ME21N para impedir que a produção pare por 5 dias por falta de assinatura.
    3. **Plano de Redução de Avarias:** Implementar manutenção preventiva na frota logística para reduzir o tempo médio entre falhas (MTBF).
    """)

st.markdown("---")

# --- BOTÕES DE DOWNLOAD DOS RELATÓRIOS ---
st.subheader("📥 4. Baixar Relatórios em Três Formatos")

col1, col2, col3 = st.columns(3)

with col1:
    buf_word = gerar_word(df, erp_escolhido)
    st.download_button(
        label="📄 Baixar Relatório Word (.docx)",
        data=buf_word,
        file_name=f"Relatorio_{erp_escolhido}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

with col2:
    buf_excel = gerar_excel(df, erp_escolhido)
    st.download_button(
        label="📊 Baixar Planilha Excel (.xlsx)",
        data=buf_excel,
        file_name=f"Gabarito_{erp_escolhido}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

with col3:
    csv_bytes = df.to_csv(index=False, sep=';').encode('utf-8-sig')
    st.download_button(
        label="📁 Baixar Arquivo CSV (.csv)",
        data=csv_bytes,
        file_name=f"Dados_Tratados_{erp_escolhido}.csv",
        mime="text/csv"
    )