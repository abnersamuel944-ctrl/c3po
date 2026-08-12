import streamlit as st
import pandas as pd
from google import genai
from google.genai import types
import os

st.set_page_config(page_title="Assistente Manthan MWS - Logística", layout="wide")

# --- 1. GERENCIAMENTO SEGURO DA CHAVE DE API ---
api_key = os.environ.get("GEMINI_API_KEY") or st.secrets.get("GEMINI_API_KEY")

if not api_key:
    st.error("🔑 Chave da API do Gemini não encontrada! Configure a variável GEMINI_API_KEY.")
    st.stop()

client = genai.Client(api_key=api_key)

# Reset de estado ao trocar de módulo
def resetar_jornada():
    if 'passos_guia' in st.session_state:
        del st.session_state.passos_guia
        del st.session_state.passo_atual

st.title("📦 Assistente Especialista em Manthan MWS (Logística)")

col_left, col_right = st.columns([1, 2])

with col_left:
    modulo_mws = st.selectbox(
        "Qual módulo do Manthan MWS você quer operar hoje?",
        [
            "Manthan MWS - Inbound (Recebimento & Guarda/Putaway)",
            "Manthan MWS - Outbound (Onda de Picking & Expedição)",
            "Manthan MWS - Gestão de Estoque & Inventário Cíclico",
            "Manthan MWS - Gestão de Pátio & Agendamento de Docas"
        ],
        on_change=resetar_jornada
    )
    
    arquivo_uploaded = st.file_uploader("Envie a planilha de operações (Excel ou CSV):", type=["xlsx", "csv"])

# --- 2. DICIONÁRIO DE LOGÍSTICA E WMS ---
DICIONARIO_WMS = {
    "sku": ["SKU", "Codigo_Produto", "Item", "Material", "Cod_Item"],
    "quantidade": ["Qtd", "Quantidade", "Volume", "Caixas", "Unidades", "Qtd_Esperada"],
    "endereco": ["Endereco", "Posicao", "Rua", "Loc", "Localizacao", "Bin"],
    "lote": ["Lote", "Batch", "Validade", "Serial"],
    "ordem_picking": ["Onda", "Wave", "Pedido", "Ordem_Separacao", "Carga"]
}

def padronizar_planilha(df):
    renomear = {}
    for col in df.columns:
        for chave, sinonimos in DICIONARIO_WMS.items():
            if str(col).strip() in sinonimos:
                renomear[col] = chave
    return df.rename(columns=renomear)

# --- 3. PROCESSAMENTO COM O GEMINI (ESPECIALISTA MANTHAN MWS) ---
if arquivo_uploaded is not None:
    try:
        if arquivo_uploaded.name.endswith('.csv'):
            df = pd.read_csv(arquivo_uploaded)
        else:
            df = pd.read_excel(arquivo_uploaded)
        
        df_padronizado = padronizar_planilha(df)
        st.success("Planilha de logística carregada e mapeada para a estrutura do Manthan MWS!")
        
        resumo_dados = f"Total de itens/registros: {len(df_padronizado)}\n"
        resumo_dados += f"Colunas identificadas: {list(df_padronizado.columns)}\n"
        resumo_dados += f"Amostra dos dados operacionais:\n{df_padronizado.head(5).to_string()}"

        system_instruction = f"""
        Você é um Especialista Sênior no sistema de WMS 'Manthan MWS' focado em Logística e Operações de Armazém.
        Sua personalidade é extremamente prática, didática, amigável e focada na rotina de armazém/centro de distribuição (CD).
        
        Módulo Atual Selecionado: {modulo_mws}
        
        Sua missão:
        1. Analisar os dados fornecidos na planilha (SKUs, quantidades, endereços, lotes, ordens de picking/recebimento).
        2. Orientar o operador/analista passo a passo de como executar essa demanda no Manthan MWS.
        3. Indicar menus específicos do MWS, uso de coletores RF (Rádio Frequência) quando aplicável, telas de liberação de ondas ou impressão de etiquetas de palete (LPN).
        4. Citar os dados reais da planilha para o usuário preencher em cada etapa.
        
        Retorne a resposta dividida EXATAMENTE em passos numerados no seguinte formato de marcação:
        ---PASSO---
        [Título do Passo | Tela/Menu do Manthan MWS]
        [Sua instrução explicando o procedimento técnico de forma clara e os dados específicos a utilizar]
        """
        
        if 'passos_guia' not in st.session_state:
            with st.spinner("O Especialista Manthan MWS está analisando sua operação..."):
                prompt = f"Aqui está o resumo da operação logística recebida:\n{resumo_dados}\n\nCrie o roteiro passo a passo para processar esta carga no Manthan MWS."
                
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=system_instruction,
                        temperature=0.3
                    )
                )
                
                raw_text = response.text
                passos = [p.strip() for p in raw_text.split("---PASSO---") if p.strip()]
                
                if passos:
                    st.session_state.passos_guia = passos
                    st.session_state.passo_atual = 0
                else:
                    st.warning("O modelo não retornou os passos no formato esperado. Tente reenviar o arquivo.")

    except Exception as e:
        st.error(f"Ocorreu um erro ao processar o arquivo de logística: {e}")

    # --- 4. PAINEL DE NAVEGAÇÃO INTERATIVO ---
    if 'passos_guia' in st.session_state and st.session_state.passos_guia:
        passos = st.session_state.passos_guia
        total_passos = len(passos)
        atual = st.session_state.passo_atual
        
        st.markdown("---")
        
        col_tit, col_btn = st.columns([3, 1])
        with col_tit:
            st.subheader(f"🏗️ Roteiro Operacional - Manthan MWS ({atual + 1} de {total_passos})")
        with col_btn:
            if st.button("🔄 Reiniciar Roteiro"):
                resetar_jornada()
                st.rerun()

        conteudo_passo = passos[atual]
        st.info(f"👷 **Especialista Manthan MWS orienta:**\n\n{conteudo_passo}")
        
        col_voltar, col_espaco, col_proximo = st.columns([1, 4, 1])
        
        with col_voltar:
            if st.button("⬅️ Passo Anterior", disabled=(atual == 0)):
                st.session_state.passo_atual -= 1
                st.rerun()
                
        with col_proximo:
            if st.button("Próximo Passo ➡️", disabled=(atual == total_passos - 1)):
                st.session_state.passo_atual += 1
                st.rerun()import os
import time
import json
import pandas as pd
from google import genai
from google.genai import types

# --- CONFIGURAÇÃO DE PASTAS ---
PASTA_ENTRADA = "entrada_dados"
PASTA_SAIDA = "saida_planilhas"
PASTA_PROCESSADOS = "historico_processados"

for pasta in [PASTA_ENTRADA, PASTA_SAIDA, PASTA_PROCESSADOS]:
    os.makedirs(pasta, exist_ok=True)

# --- CHAVE DE API GEMINI ---
api_key = os.environ.get("GEMINI_API_KEY")
if not api_key:
    print("❌ ERRO: Variável de ambiente GEMINI_API_KEY não encontrada!")
    print("Defina a chave no terminal: export GEMINI_API_KEY='sua_chave'")
    exit(1)

client = genai.Client(api_key=api_key)

SYSTEM_PROMPT = """
Você é um motor de extração de dados para sistemas ERP e WMS (Manthan MWS, SAP, Protheus).
Sua tarefa é ler mensagens de texto informais (como conversas de WhatsApp, e-mails ou anotações) e extrair os dados em formato de tabela JSON limpa.

Mapeamento obrigatório de colunas para o ERP:
- SKU (ou Codigo_Produto)
- Item (Descrição do produto)
- Quantidade (Apenas o número)
- Endereco (Posição/Rua no armazém)
- Lote (Número de lote/batch)
- Ordem (Número da OP, Pedido ou Carga)
- Data (Data limite/prazo)

Retorne EXATAMENTE um JSON puro no formato de lista de objetos.
Exemplo de retorno:
[
  {"SKU": "ML-04", "Item": "Molas de Aço", "Quantidade": 500, "Endereco": "RUA-A-01-2", "Lote": "L2026-A", "Ordem": "99402", "Data": "15/08/2026"}
]
Não inclua nenhuma explicação, apenas o código JSON.
"""

def processar_texto_com_gemini(texto_raw):
    """Envia o texto cru do WhatsApp para a IA estruturar nas colunas do ERP."""
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=f"Extraia os dados operacionais deste texto:\n\n{texto_raw}",
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=0.1
        )
    )
    
    # Limpa possíveis marcações de código markdown do retorno
    raw_response = response.text.strip()
    if raw_response.startswith("```json"):
        raw_response = raw_response[7:-3].strip()
    elif raw_response.startswith("```"):
        raw_response = raw_response[3:-3].strip()
        
    return json.loads(raw_response)

def executar_robo():
    print("🤖 Robô ERP/MWS Iniciado!")
    print(f"📂 Vigia ativado na pasta: ./{PASTA_ENTRADA}")
    print("Aguardando novos arquivos para converter...\n")

    while True:
        arquivos = [f for f in os.listdir(PASTA_ENTRADA) if not f.startswith('.')]
        
        for nome_arquivo in arquivos:
            caminho_arquivo = os.path.join(PASTA_ENTRADA, nome_arquivo)
            print(f"⚡ Novo arquivo detectado: {nome_arquivo}")
            
            try:
                # 1. Leitura do arquivo colado na pasta
                with open(caminho_arquivo, "r", encoding="utf-8") as f:
                    conteudo = f.read()

                print("🧠 Convertendo texto informal na linguagem do ERP via Gemini...")
                dados_json = processar_texto_com_gemini(conteudo)

                # 2. Converte JSON para DataFrame do Pandas
                df = pd.DataFrame(dados_json)

                # 3. Gera a planilha Excel pronta para o ERP
                timestamp = time.strftime("%Y%m%d_%H%M%S")
                nome_excel = f"Planilha_ERP_MWS_{timestamp}.xlsx"
                caminho_excel = os.path.join(PASTA_SAIDA, nome_excel)

                df.to_excel(caminho_excel, index=False, engine='openpyxl')
                print(f"✅ PLANILHA GERADA COM SUCESSO: {caminho_excel}")

                # 4. Move o arquivo original para histórico
                caminho_historico = os.path.join(PASTA_PROCESSADOS, f"{timestamp}_{nome_arquivo}")
                os.rename(caminho_arquivo, caminho_historico)
                print(f"📦 Arquivo fonte movido para historico.\n")

            except Exception as e:
                print(f"❌ Erro ao processar arquivo {nome_arquivo}: {e}\n")

        # Aguarda 3 segundos antes de checar a pasta novamente
        time.sleep(3)

if __name__ == "__main__":
    executar_robo()import os
import io
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import tkinter as tk
from tkinter import filedialog, messagebox

def aplicar_sombra_tabela(cell, color_hex):
    """Aplica cor de fundo a uma célula da tabela do Word."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def obter_pasta_meu_drive():
    """Identifica o caminho para a pasta 'Meu Drive' ou retorna a pasta do usuário."""
    candidatos = [
        os.path.expanduser("~/Google Drive"),
        os.path.expanduser("~/Meu Drive"),
        os.path.expanduser("~/GoogleDrive"),
        os.path.expanduser("~")
    ]
    for caminho in candidatos:
        if os.path.exists(caminho):
            return caminho
    return "."

def selecionar_arquivo():
    """Caixa de entrada visual para seleção de arquivo CSV."""
    root = tk.Tk()
    root.withdraw() # Oculta a janela principal do Tkinter
    root.attributes('-topmost', True) # Traz a janela para frente

    messagebox.showinfo("Robô de Relatórios", "Por favor, selecione o arquivo CSV de entrada.")
    
    caminho_csv = filedialog.askopenfilename(
        title="Selecione o arquivo CSV de dados",
        filetypes=[("Arquivos CSV", "*.csv"), ("Todos os arquivos", "*.*")]
    )
    
    return caminho_csv

def gerar_relatorios():
    # 1. Caixa de entrada para seleção do arquivo
    caminho_csv_input = selecionar_arquivo()
    
    if not caminho_csv_input:
        print("❌ Nenhum arquivo selecionado. Processo cancelado.")
        return

    # 2. Leitura do arquivo selecionado
    try:
        # Tenta ler separador ; ou ,
        try:
            df = pd.read_csv(caminho_csv_input, sep=';')
            if len(df.columns) <= 1:
                df = pd.read_csv(caminho_csv_input, sep=',')
        except Exception:
            df = pd.read_csv(caminho_csv_input, sep=',')
    except Exception as e:
        messagebox.showerror("Erro", f"Não foi possível ler o arquivo CSV:\n{e}")
        return

    # 3. Definição da Pasta de Destino (Meu Drive)
    pasta_destino = obter_pasta_meu_drive()
    print(f"📁 Salvando arquivos em: {pasta_destino}")

    caminho_docx = os.path.join(pasta_destino, "Relatorio_Diagnostico_Operacional.docx")
    caminho_excel = os.path.join(pasta_destino, "Relatorio_Diagnostico_Operacional.xlsx")
    caminho_csv_out = os.path.join(pasta_destino, "Relatorio_Diagnostico_Operacional.csv")

    # 4. Cálculo dos KPIs e Métricas
    custo_total = df['Custo_Parada_R$'].sum() if 'Custo_Parada_R$' in df.columns else 0
    media_atraso = df['Dias_Atraso'].mean() if 'Dias_Atraso' in df.columns else 0
    total_linhas = len(df)
    paradas = len(df[df['Status_Linha'] == 'PARADA']) if 'Status_Linha' in df.columns else 0
    lentas = len(df[df['Status_Linha'] == 'LENTA']) if 'Status_Linha' in df.columns else 0
    em_transito = len(df[df['Status_Linha'] == 'EM TRÂNSITO']) if 'Status_Linha' in df.columns else 0

    # -------------------------------------------------------------
    # 5. GERAR ARQUIVO EXCEL (.xlsx)
    # -------------------------------------------------------------
    with pd.ExcelWriter(caminho_excel, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Dados Gargalos', index=False)
        
        # Cria aba de resumo/KPIs
        df_kpis = pd.DataFrame({
            'Indicador': ['Custo Total Impacto (R$)', 'Média Dias Atraso', 'Total Postos Afetados', 'Linhas Paradas', 'Linhas Lentas', 'Em Trânsito'],
            'Valor': [custo_total, media_atraso, total_linhas, paradas, lentas, em_transito]
        })
        df_kpis.to_excel(writer, sheet_name='Resumo KPIs', index=False)

    # -------------------------------------------------------------
    # 6. GERAR ARQUIVO CSV (.csv)
    # -------------------------------------------------------------
    df.to_csv(caminho_csv_out, sep=';', index=False, encoding='utf-8-sig')

    # -------------------------------------------------------------
    # 7. GERAR DOCUMENTO WORD (.docx)
    # -------------------------------------------------------------
    doc = Document()
    doc.add_heading('📊 Relatório de Diagnóstico e Impacto Operacional', level=0)
    
    # Seção 1
    doc.add_heading('1. Do que se trata este documento?', level=2)
    doc.add_paragraph(
        "Este documento é um relatório de alertas/avisos de paradas e gargalos na operação industrial e logística. "
        "Ele mapeia falhas críticas que estão afetando o fluxo de produção e transporte, identificando o local (centro de trabalho), "
        "o motivo da desaceleração/parada, o tempo de atraso gerado e o prejuízo financeiro direto associado a cada ocorrência."
    )

    # Seção 2: Tabela
    doc.add_heading('2. Resumo dos Dados (Tabela Tratada)', level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name.replace('_', ' ')
        aplicar_sombra_tabela(hdr_cells[i], "1F4E79")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float) and 'Custo' in df.columns[i]:
                row_cells[i].text = f"R$ {val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            else:
                row_cells[i].text = str(val)

    doc.add_paragraph("")

    # Seção 3: KPIs
    doc.add_heading('3. Indicadores-Chave (KPIs)', level=2)
    p_kpis = doc.add_paragraph()
    p_kpis.add_run("• Custo Total de Impacto Financeiro: ").bold = True
    p_kpis.add_run(f"R$ {custo_total:,.2f}\n".replace(',', 'X').replace('.', ',').replace('X', '.'))
    p_kpis.add_run("• Média de Dias de Atraso: ").bold = True
    p_kpis.add_run(f"{media_atraso:.2f} dias\n".replace('.', ','))
    p_kpis.add_run("• Total de Linhas/Rotas Afetadas: ").bold = True
    p_kpis.add_run(f"{total_linhas} postos de trabalho\n")
    p_kpis.add_run("• Linhas Totalmente Paradas: ").bold = True
    p_kpis.add_run(f"{paradas} (Prensa-01 e Montagem-01)\n")
    p_kpis.add_run("• Linha com Operação Parcial/Lenta: ").bold = True
    p_kpis.add_run(f"{lentas} (Solda-02)\n")
    p_kpis.add_run("• Frota/Rota Com Problema Logístico: ").bold = True
    p_kpis.add_run(f"{em_transito} (Rota-Norte)")

    # Seção 4: Análise Crítica
    doc.add_heading('4. Análise Crítica dos Problemas Detectados', level=2)
    analises = [
        ("1. Gargalo Administrativo/Sistemas (Maior Impacto Financeiro e de Tempo):", 
         "Ocorrência: AV-1004 (Montagem-01)\nProblema: Linha parada aguardando liberação do Pedido ME21N no SAP.\nImpacto: Prejuízo de R$ 22.100,00."),
        ("2. Gargalo de Suprimentos/Estoque:", 
         "Ocorrência: AV-1001 (Prensa-01)\nProblema: Falta de Chapa de Aço (4,5 dias de parada).\nImpacto: Prejuízo de R$ 18.500,00."),
        ("3. Gargalo Logístico / Manutenção de Frota:", 
         "Ocorrência: AV-1003 (Rota-Norte)\nProblema: Frota quebrada gerando 3 dias de atraso.\nImpacto: Prejuízo de R$ 9.400,00."),
        ("4. Gargalo de Capacidade Produtiva:", 
         "Ocorrência: AV-1002 (Solda-02)\nProblema: Sobrecarga de ordens de produção.\nImpacto: Operação lenta com custo de R$ 6.200,00.")
    ]
    for tit, txt in analises:
        p = doc.add_paragraph()
        p.add_run(f"{tit}\n").bold = True
        p.add_run(txt)

    # Seção 5: Plano de Ação
    doc.add_heading('5. Plano de Ação Recomendado (Próximos Passos)', level=2)
    acoes = [
        ("🔴 Prioridade 1 (Imediata) - Montagem-01: ", "Entrar em contato com Compras/PCP para liberar o pedido ME21N."),
        ("🔴 Prioridade 2 (Imediata) - Prensa-01: ", "Agilizar recebimento emergencial de Chapas de Aço."),
        ("🟡 Prioridade 3 - Rota-Norte: ", "Acionar veículo reserva para transbordo da carga."),
        ("🟢 Prioridade 4 - Solda-02: ", "Redistribuir ordens de serviço ou aplicar horas extras.")
    ]
    for prio, desc in acoes:
        p = doc.add_paragraph()
        p.add_run(prio).bold = True
        p.add_run(desc)

    doc.save(caminho_docx)

    # Mensagem de sucesso
    mensagem_sucesso = (
        f"✅ Relatórios gerados com sucesso na pasta Meu Drive!\n\n"
        f"📄 Word: {caminho_docx}\n"
        f"📊 Excel: {caminho_excel}\n"
        f"📁 CSV: {caminho_csv_out}"
    )
    print(mensagem_sucesso)
    messagebox.showinfo("Sucesso!", mensagem_sucesso)

if __name__ == "__main__":
    gerar_relatorios()