import os
import io
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import tkinter as tk
from tkinter import filedialog, messagebox

def aplicar_sombra_tabela(cell, color_hex):
    """Aplica cor de fundo a uma célula da tabela do Word."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def obter_pasta_meu_drive():
    """Identifica o caminho para a pasta 'Meu Drive' ou retorna a pasta do usuário."""
    candidatos = [
        os.path.expanduser("~/Google Drive"),
        os.path.expanduser("~/Meu Drive"),
        os.path.expanduser("~/GoogleDrive"),
        os.path.expanduser("~")
    ]
    for caminho in candidatos:
        if os.path.exists(caminho):
            return caminho
    return "."

def selecionar_arquivo():
    """Caixa de entrada visual para seleção de arquivo CSV."""
    root = tk.Tk()
    root.withdraw() # Oculta a janela principal do Tkinter
    root.attributes('-topmost', True) # Traz a janela para frente

    messagebox.showinfo("Robô de Relatórios", "Por favor, selecione o arquivo CSV de entrada.")
    
    caminho_csv = filedialog.askopenfilename(
        title="Selecione o arquivo CSV de dados",
        filetypes=[("Arquivos CSV", "*.csv"), ("Todos os arquivos", "*.*")]
    )
    
    return caminho_csv

def gerar_relatorios():
    # 1. Caixa de entrada para seleção do arquivo
    caminho_csv_input = selecionar_arquivo()
    
    if not caminho_csv_input:
        print("❌ Nenhum arquivo selecionado. Processo cancelado.")
        return

    # 2. Leitura do arquivo selecionado
    try:
        # Tenta ler separador ; ou ,
        try:
            df = pd.read_csv(caminho_csv_input, sep=';')
            if len(df.columns) <= 1:
                df = pd.read_csv(caminho_csv_input, sep=',')
        except Exception:
            df = pd.read_csv(caminho_csv_input, sep=',')
    except Exception as e:
        messagebox.showerror("Erro", f"Não foi possível ler o arquivo CSV:\n{e}")
        return

    # 3. Definição da Pasta de Destino (Meu Drive)
    pasta_destino = obter_pasta_meu_drive()
    print(f"📁 Salvando arquivos em: {pasta_destino}")

    caminho_docx = os.path.join(pasta_destino, "Relatorio_Diagnostico_Operacional.docx")
    caminho_excel = os.path.join(pasta_destino, "Relatorio_Diagnostico_Operacional.xlsx")
    caminho_csv_out = os.path.join(pasta_destino, "Relatorio_Diagnostico_Operacional.csv")

    # 4. Cálculo dos KPIs e Métricas
    custo_total = df['Custo_Parada_R$'].sum() if 'Custo_Parada_R$' in df.columns else 0
    media_atraso = df['Dias_Atraso'].mean() if 'Dias_Atraso' in df.columns else 0
    total_linhas = len(df)
    paradas = len(df[df['Status_Linha'] == 'PARADA']) if 'Status_Linha' in df.columns else 0
    lentas = len(df[df['Status_Linha'] == 'LENTA']) if 'Status_Linha' in df.columns else 0
    em_transito = len(df[df['Status_Linha'] == 'EM TRÂNSITO']) if 'Status_Linha' in df.columns else 0

    # -------------------------------------------------------------
    # 5. GERAR ARQUIVO EXCEL (.xlsx)
    # -------------------------------------------------------------
    with pd.ExcelWriter(caminho_excel, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Dados Gargalos', index=False)
        
        # Cria aba de resumo/KPIs
        df_kpis = pd.DataFrame({
            'Indicador': ['Custo Total Impacto (R$)', 'Média Dias Atraso', 'Total Postos Afetados', 'Linhas Paradas', 'Linhas Lentas', 'Em Trânsito'],
            'Valor': [custo_total, media_atraso, total_linhas, paradas, lentas, em_transito]
        })
        df_kpis.to_excel(writer, sheet_name='Resumo KPIs', index=False)

    # -------------------------------------------------------------
    # 6. GERAR ARQUIVO CSV (.csv)
    # -------------------------------------------------------------
    df.to_csv(caminho_csv_out, sep=';', index=False, encoding='utf-8-sig')

    # -------------------------------------------------------------
    # 7. GERAR DOCUMENTO WORD (.docx)
    # -------------------------------------------------------------
    doc = Document()
    doc.add_heading('📊 Relatório de Diagnóstico e Impacto Operacional', level=0)
    
    # Seção 1
    doc.add_heading('1. Do que se trata este documento?', level=2)
    doc.add_paragraph(
        "Este documento é um relatório de alertas/avisos de paradas e gargalos na operação industrial e logística. "
        "Ele mapeia falhas críticas que estão afetando o fluxo de produção e transporte, identificando o local (centro de trabalho), "
        "o motivo da desaceleração/parada, o tempo de atraso gerado e o prejuízo financeiro direto associado a cada ocorrência."
    )

    # Seção 2: Tabela
    doc.add_heading('2. Resumo dos Dados (Tabela Tratada)', level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name.replace('_', ' ')
        aplicar_sombra_tabela(hdr_cells[i], "1F4E79")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float) and 'Custo' in df.columns[i]:
                row_cells[i].text = f"R$ {val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            else:
                row_cells[i].text = str(val)

    doc.add_paragraph("")

    # Seção 3: KPIs
    doc.add_heading('3. Indicadores-Chave (KPIs)', level=2)
    p_kpis = doc.add_paragraph()
    p_kpis.add_run("• Custo Total de Impacto Financeiro: ").bold = True
    p_kpis.add_run(f"R$ {custo_total:,.2f}\n".replace(',', 'X').replace('.', ',').replace('X', '.'))
    p_kpis.add_run("• Média de Dias de Atraso: ").bold = True
    p_kpis.add_run(f"{media_atraso:.2f} dias\n".replace('.', ','))
    p_kpis.add_run("• Total de Linhas/Rotas Afetadas: ").bold = True
    p_kpis.add_run(f"{total_linhas} postos de trabalho\n")
    p_kpis.add_run("• Linhas Totalmente Paradas: ").bold = True
    p_kpis.add_run(f"{paradas} (Prensa-01 e Montagem-01)\n")
    p_kpis.add_run("• Linha com Operação Parcial/Lenta: ").bold = True
    p_kpis.add_run(f"{lentas} (Solda-02)\n")
    p_kpis.add_run("• Frota/Rota Com Problema Logístico: ").bold = True
    p_kpis.add_run(f"{em_transito} (Rota-Norte)")

    # Seção 4: Análise Crítica
    doc.add_heading('4. Análise Crítica dos Problemas Detectados', level=2)
    analises = [
        ("1. Gargalo Administrativo/Sistemas (Maior Impacto Financeiro e de Tempo):", 
         "Ocorrência: AV-1004 (Montagem-01)\nProblema: Linha parada aguardando liberação do Pedido ME21N no SAP.\nImpacto: Prejuízo de R$ 22.100,00."),
        ("2. Gargalo de Suprimentos/Estoque:", 
         "Ocorrência: AV-1001 (Prensa-01)\nProblema: Falta de Chapa de Aço (4,5 dias de parada).\nImpacto: Prejuízo de R$ 18.500,00."),
        ("3. Gargalo Logístico / Manutenção de Frota:", 
         "Ocorrência: AV-1003 (Rota-Norte)\nProblema: Frota quebrada gerando 3 dias de atraso.\nImpacto: Prejuízo de R$ 9.400,00."),
        ("4. Gargalo de Capacidade Produtiva:", 
         "Ocorrência: AV-1002 (Solda-02)\nProblema: Sobrecarga de ordens de produção.\nImpacto: Operação lenta com custo de R$ 6.200,00.")
    ]
    for tit, txt in analises:
        p = doc.add_paragraph()
        p.add_run(f"{tit}\n").bold = True
        p.add_run(txt)

    # Seção 5: Plano de Ação
    doc.add_heading('5. Plano de Ação Recomendado (Próximos Passos)', level=2)
    acoes = [
        ("🔴 Prioridade 1 (Imediata) - Montagem-01: ", "Entrar em contato com Compras/PCP para liberar o pedido ME21N."),
        ("🔴 Prioridade 2 (Imediata) - Prensa-01: ", "Agilizar recebimento emergencial de Chapas de Aço."),
        ("🟡 Prioridade 3 - Rota-Norte: ", "Acionar veículo reserva para transbordo da carga."),
        ("🟢 Prioridade 4 - Solda-02: ", "Redistribuir ordens de serviço ou aplicar horas extras.")
    ]
    for prio, desc in acoes:
        p = doc.add_paragraph()
        p.add_run(prio).bold = True
        p.add_run(desc)

    doc.save(caminho_docx)

    # Mensagem de sucesso
    mensagem_sucesso = (
        f"✅ Relatórios gerados com sucesso na pasta Meu Drive!\n\n"
        f"📄 Word: {caminho_docx}\n"
        f"📊 Excel: {caminho_excel}\n"
        f"📁 CSV: {caminho_csv_out}"
    )
    print(mensagem_sucesso)
    messagebox.showinfo("Sucesso!", mensagem_sucesso)

if __name__ == "__main__":
    gerar_relatorios()